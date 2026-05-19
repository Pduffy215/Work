from pydoc import text
import sys
from docx import Document
import openpyxl
from openpyxl.styles import Alignment
import time
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
import os 
import copy
import re
from datetime import datetime

BASE_DIR = os.path.dirname(os.path.abspath(sys.executable if getattr(sys, 'frozen', False) else __file__))  
LOG_FILE = os.path.join(BASE_DIR, "watcher_log.txt")              # Log file stored next to EXE/script
INCOMING_FOLDER = os.path.join(BASE_DIR, "IncomingFiles")        # Folder to watch stored next to EXE/script
EXCEL_FILE = os.path.join(BASE_DIR, "temp.xlsx") # Excel file stored next to EXE/script


# Extract underlined starred text from Word documents, counting leading stars for severity
def find_highlighted_text(doc):
    equipment_stars = {}  # Dictionary to store equipment names and their corresponding star counts
    seen = set()  # To track already processed equipment names and avoid duplicates

    for para in doc.paragraphs:
        underlined_runs = [] # Temporary list to collect underlined runs for the current paragraph
        
        for run in para.runs:
            if run.font.underline: # Check if the run is underlined
                underlined_runs.append(run.text)
            else:
                if underlined_runs:  # If we have collected underlined runs, process them before moving on
                    text = "".join(underlined_runs).strip() # Combine the collected underlined runs into a single string and strip whitespace
                    
                    if text and text not in seen: 
                        seen.add(text) # Add the equipment name to the seen set to avoid duplicates
                        
                        star_count = 0 
                        for char in text:
                            if char == "*": # Count leading stars to determine severity
                                star_count += 1
                            else:
                                break
                        
                        clean_text = text.lstrip("*").strip()
                        equipment_stars[clean_text] = star_count  # Store the cleaned equipment name and its corresponding star count in the dictionary           
                    underlined_runs = []

        if underlined_runs:
            text = "".join(underlined_runs).strip() # Combine any remaining underlined runs at the end of the paragraph
            
            if text and text not in seen:
                seen.add(text)
                
                star_count = 0
                for char in text:
                    if char == "*":
                        star_count += 1
                    else:
                        break
                
                clean_text = text.lstrip("*").strip()
                equipment_stars[clean_text] = star_count

    return equipment_stars

# Extract the second date from the Word document in the format "Month Day, Year"
def find_date(doc):
    pattern = r"[A-Za-z]+\s+\d{1,2},\s+\d{4}"

    all_matches = []

    for para in doc.paragraphs:
        all_matches.extend(re.findall(pattern, para.text))

    return all_matches[1] if len(all_matches) >= 2 else None



# Write the extracted equipment names and their star counts to the Excel file, applying color coding based on severity
def write_to_excel(data, date, plant_name, filename):
    color_map = {
    1: "99CCFF",  # Blue
    2: "FFFF00",  # Yellow
    3: "FFA500",  # Orange
    4: "FF3300",  # Red
    0: "008000"   # Green
}
    
    wb = openpyxl.load_workbook(filename) # Load the existing Excel file
    sheet = wb.active 

    for key, value in data.items(): # Iterate through the extracted equipment names and their star counts
        search_text = key

        # Iterate through the rows in the Excel sheet, starting from row 2 and checking columns B and C
        for row in sheet.iter_rows(min_row=2, min_col=1, max_col=4, values_only=False): 
            cell_plant = row[0]

            if not cell_plant.value or cell_plant.value.strip() == "":
                continue

            if cell_plant.value.split()[0] != plant_name.upper():
                continue                     # Skip rows that do not match the plant name
            
            cell_equipment = row[1]
            cell_stars = row[2]  
            cell_date = row[3]
                      # star count should be written in column C 

 
            if cell_equipment.value.replace(" ", "") == search_text.upper().replace(" ", ""): # If the equipment name matches the search text, write the star count and apply color coding
                i = 0
                stars = ""  
                if cell_stars.value != None:
                    print(f"key {search_text} Already filled")
                    continue
                
                cell_date.value = date # Update the date in column D to the current date

                while i < value: # Append the appropriate number of stars based on the severity count
                    stars += "*"
                    i = i + 1
                cell_stars.value = stars
                star_count = str(cell_stars.value or "").count("*") # Count the number of stars to determine the severity level for color coding
                cell_stars.fill = openpyxl.styles.PatternFill(start_color=color_map.get(star_count, "FFFFFF"), end_color=color_map.get(star_count, "FFFFFF"), fill_type="solid")
                break

    wb.save(filename)

# Wait for the file to be fully copied and ready to be opened
def wait_for_file(file_path, timeout=10): 
    start_time = time.time()
    while True:
        try:
            doc = Document(file_path)
            return doc  
        except Exception:
            if time.time() - start_time > timeout:
                raise TimeoutError(f"File not ready: {file_path}")
            time.sleep(0.5)


# Watchdog event handler to monitor the folder for new files and process them accordingly
class MyHandler(FileSystemEventHandler):
    def on_created(self, event):

        if event.is_directory:
            return
        
        file_path = event.src_path

        if not file_path.lower().endswith(".docx"):
            return

        file_name = os.path.splitext(os.path.basename(file_path))[0] # Extract the plant name from the file name
        plant_name = file_name.split()[1]


        # Ignore temporary Word files
        if os.path.basename(file_path).startswith("~$"):
            return

        print(f"New file detected: {file_path}")

        with open(LOG_FILE, "a") as f:
            f.write(f"{time.ctime()} - New file detected: {file_path}\n")

        try:
            # Wait for file to finish copying
            doc = wait_for_file(file_path)

            # Extract date
            date = find_date(doc)
            dt = datetime.strptime(date, "%B %d, %Y")
            formatted_date = dt.strftime("%#m/%#d/%Y")  

            # Extract stars
            equipment_stars = find_highlighted_text(doc)

            with open(LOG_FILE, "a") as f:
                f.write(f"{time.ctime()} - Extracted data: {equipment_stars}\n")

            # Write to Excel
            write_to_excel(equipment_stars, formatted_date, plant_name, EXCEL_FILE)

            with open(LOG_FILE, "a") as f:
                f.write(f"{time.ctime()} - Excel updated successfully\n")

        except Exception as e:

            print(f"ERROR: {e}")

            with open(LOG_FILE, "a") as f:
                f.write(f"{time.ctime()} - ERROR: {e}\n")


def swap_cells(cell1, cell2):

    # store originals
    v1, v2 = cell1.value, cell2.value
    f1, f2 = copy.copy(cell1.fill), copy.copy(cell2.fill)
    font1, font2 = copy.copy(cell1.font), copy.copy(cell2.font)
    b1, b2 = copy.copy(cell1.border), copy.copy(cell2.border)
    a1, a2 = copy.copy(cell1.alignment), copy.copy(cell2.alignment)
    n1, n2 = cell1.number_format, cell2.number_format

    # swap values
    cell1.value, cell2.value = v2, v1

    # swap styles
    cell1.fill, cell2.fill = f2, f1
    cell1.font, cell2.font = font2, font1
    cell1.border, cell2.border = b2, b1
    cell1.alignment, cell2.alignment = a2, a1
    cell1.number_format = n2
    cell2.number_format = n1

def swap_columns(sheet, col1, col2):
    for r in range(1, sheet.max_row + 1):
        swap_cells(
            sheet.cell(row=r, column=col1),
            sheet.cell(row=r, column=col2)
        )

def copy_column_format(sheet, source_col, target_col):
    max_row = sheet.max_row

    for r in range(1, max_row + 1):
        src = sheet.cell(row=r, column=source_col)
        dst = sheet.cell(row=r, column=target_col)

        dst.font = copy.copy(src.font)
        dst.fill = copy.copy(src.fill)
        dst.border = copy.copy(src.border)
        dst.number_format = src.number_format
        dst.alignment = Alignment(
            wrap_text=True,
            horizontal=src.alignment.horizontal,
            vertical=src.alignment.vertical,
            shrink_to_fit=src.alignment.shrink_to_fit
        )

def add_column(file):
    wb = openpyxl.load_workbook(file)
    sheet = wb.active

    new_col = input("Enter new column name: ")

    sheet.insert_cols(3)  # Insert a new column at index 3 (C)
    copy_column_format(sheet, 2, 3)  # Copy formatting from column B to the new column C
    sheet.cell(row=1, column=3).value = new_col
    swap_columns(sheet, 4, 5)
    wb.save(file)
    


def main():
    with open(LOG_FILE, "a") as f:
        f.write(f"{time.ctime()} - Started watcher\n") # Log the start of the watcher


    folder_to_watch = INCOMING_FOLDER # Set the folder to watch for new files
    event_handler = MyHandler()
    observer = Observer()
    observer.schedule(event_handler, folder_to_watch, recursive=True) # Schedule the event handler to monitor the specified folder and its subfolders
    observer.start()


    print(f"Watching folder: {folder_to_watch}") 
    try:
        while True: # Keep the script running to continuously monitor the folder for new files
            choice = input("Do you want to add a column to the Excel file? (yes/no): \n").strip().lower()
            if choice == "yes":
                print("Adding a new column to the Excel file...")
                add_column(EXCEL_FILE)
            time.sleep(1)
    except KeyboardInterrupt:
        observer.stop() 
    observer.join() 


if __name__ == "__main__":
    main()


